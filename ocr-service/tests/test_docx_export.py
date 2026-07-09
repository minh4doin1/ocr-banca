"""Tests xuất Word (.docx) từ kết quả OCR."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from app.models.schemas import CellData, OcrResult, PageResult, TableData
from app.services.docx_service import export_to_docx


def _sample_result() -> OcrResult:
    cells = [
        CellData(row=0, col=0, text="1", confidence=0.95),
        CellData(row=0, col=1, text="Pham Anh Tuan", confidence=0.92),
        CellData(row=0, col=2, text="3526", confidence=0.88),
        CellData(row=0, col=3, text="CN Tien Giang", confidence=0.9),
        CellData(row=0, col=4, text="QSOPTUAN", confidence=0.99),
        CellData(row=0, col=5, text="038093031725", confidence=0.97),
        CellData(row=0, col=6, text="tuan@agribank.com.vn", confidence=0.99),
        CellData(row=0, col=7, text="0987654321", confidence=0.95),
        CellData(row=0, col=8, text="dai ly vien", confidence=0.9),
        CellData(row=0, col=9, text="", confidence=1.0),
    ]
    table = TableData(
        table_index=0,
        num_rows=1,
        num_cols=10,
        cells=cells,
        html="",
        table_kind="sso_agribank",
    )
    now = datetime.now()
    return OcrResult(
        job_id="job-docx-export",
        filename="danh_sach.pdf",
        total_pages=1,
        pages=[
            PageResult(
                page_number=1,
                image_path="",
                tables=[table],
                raw_text="",
            )
        ],
        is_complete=True,
        created_at=now,
        updated_at=now,
    )


def test_export_to_docx_creates_styled_file(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "storage_dir", str(tmp_path))
    result = _sample_result()
    out = export_to_docx(result)
    assert out.exists()
    assert out.suffix == ".docx"

    doc = Document(str(out))
    assert doc.tables
    table = doc.tables[0]
    assert len(table.rows) >= 2
    assert table.cell(0, 0).text == "STT"
    assert table.cell(1, 4).text == "QSOPTUAN"
    # Tiêu đề tài liệu
    assert any("Danh sách user SSO" in p.text for p in doc.paragraphs)
