"""Tests cho import_from_docx — đọc bảng Word (.docx) trực tiếp."""

from __future__ import annotations

from pathlib import Path

import pytest

from app.services.docx_service import import_from_docx


def _make_docx(path: Path, rows: list[list[str]]) -> None:
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row_vals in enumerate(rows):
        for c, val in enumerate(row_vals):
            table.cell(r, c).text = val
    doc.save(str(path))


def test_import_from_docx_reads_table(tmp_path):
    docx_path = tmp_path / "danh_sach.docx"
    _make_docx(
        docx_path,
        [
            ["STT", "Ho va ten", "Phong/Don vi", "User IPCAS", "So CCCD",
             "Email tai Agribank", "So dien thoai", "Phan quyen", "Ma DV"],
            ["1", "Nguyen Van A", "CN 6900", "HQPT01", "001234567890",
             "a@agribank.com.vn", "0982867163", "dai ly vien", "82204001"],
            ["2", "Tran Thi B", "CN 6900", "HQPT02", "001234567891",
             "b@agribank.com.vn", "0982867164", "ke toan vien", "82204001"],
        ],
    )

    result = import_from_docx(docx_path, "job-docx", "danh_sach.docx")

    assert result.total_pages == 1
    assert result.is_complete is True
    page = result.pages[0]
    assert len(page.tables) == 1
    table = page.tables[0]
    # Header bị loại bỏ -> còn 2 dòng dữ liệu.
    assert table.num_rows == 2
    assert table.table_kind == "sso_agribank"
    grid = {(c.row, c.col): c.text for c in table.cells}
    assert grid[(0, 1)] == "Nguyen Van A"
    assert grid[(1, 1)] == "Tran Thi B"
    # Đọc trực tiếp -> confidence tuyệt đối.
    assert all(c.confidence == 1.0 for c in table.cells)


def test_import_from_docx_multiline_cell(tmp_path):
    docx_path = tmp_path / "multiline.docx"
    _make_docx(
        docx_path,
        [
            ["STT", "Ho va ten", "Phong/Don vi", "User IPCAS", "So CCCD",
             "Email tai Agribank"],
            ["1", "Nguyen\nVan A", "CN 6900", "HQPT01", "001234567890",
             "a@agribank.com.vn"],
        ],
    )

    result = import_from_docx(docx_path, "job-docx", "multiline.docx")
    grid = {(c.row, c.col): c.text for c in result.pages[0].tables[0].cells}
    assert grid[(0, 1)] == "Nguyen Van A"


def test_map_table_accent_header_and_annotation_row():
    """Tiêu đề có dấu + dòng (1)(2)… vẫn nhận được cột email."""
    from app.models.schemas import CellData, TableData
    from app.services.user_mapping import map_table_to_users

    matrix = [
        [
            "STT", "Họ và tên", "Mã chi nhánh", "Tên Chi nhánh", "User IPCAS",
            "Số CCCD", "Email tại Agribank", "SĐT",
            "Phân quyền (Quản trị/Đại lý viên/Kế toán viên/Phê duyệt viên)",
            "Mã liên ngân hàng",
        ],
        ["(1)", "(2)", "(3)", "(4)", "(5)", "(6)", "(7)", "(8)", "(9)", "(10)"],
        [
            "1", "Lê Mai Phương", "6900", "Hội sở Agribank CN Tiền Giang", "TGILMP",
            "001234567890", "phuonglemail@agribank.com.vn", "0589.636.987",
            "Đại lý viên", "82204001",
        ],
    ]
    cells = [
        CellData(row=r, col=c, text=txt, confidence=1.0, bbox=[])
        for r, row in enumerate(matrix)
        for c, txt in enumerate(row)
    ]
    table = TableData(
        table_index=0, num_rows=len(matrix), num_cols=10,
        cells=cells, html="", table_kind="",
    )
    users, warnings = map_table_to_users(table)
    assert not warnings
    assert len(users) == 1
    assert users[0].phone == "0589636987"
    assert users[0].email.endswith("@agribank.com.vn")


def test_normalize_phone_strips_dots():
    from app.services.user_mapping import _normalize_phone

    assert _normalize_phone("0589.636.987") == "0589636987"
    assert _normalize_phone("0982 867 163") == "0982867163"


def test_import_from_docx_reads_table_10_columns(tmp_path):
    """Mẫu SSO mới 10 cột (Mẫu 01/SSO)."""
    docx_path = tmp_path / "sso_moi.docx"
    _make_docx(
        docx_path,
        [
            [
                "STT", "Ho va ten", "Ma chi nhanh", "Ten Chi nhanh", "User IPCAS",
                "So CCCD", "Email tai Agribank", "SDT", "Phan quyen", "Ma lien ngan hang",
            ],
            [
                "1", "Le Mai Phuong", "6900", "Hoi so Agribank CN Tien Giang", "TGILMP",
                "001234567890", "phuonglemail@agribank.com.vn", "0589636987",
                "Dai ly vien", "82204001",
            ],
        ],
    )

    result = import_from_docx(docx_path, "job-10", "sso_moi.docx")
    table = result.pages[0].tables[0]
    assert table.num_cols == 10
    assert table.num_rows == 1

    from app.services.user_mapping import map_table_to_users

    users, warnings = map_table_to_users(table)
    assert not warnings
    assert len(users) == 1
    assert users[0].branch_code == "6900"
    assert users[0].branch_name == "Hoi so Agribank CN Tien Giang"
    assert users[0].unit_code == "82204001"


def test_map_table_with_sso_labels_header_any_order():
    """Bảng có header nhãn SSO (thứ tự cột bất kỳ) phải nhận được cột email."""
    from app.models.schemas import CellData, TableData
    from app.services.user_mapping import map_table_to_users

    header = [
        "Ho va ten", "Email tại Agribank", "User IPCAS", "So CCCD",
        "So dien thoai", "Phan quyen", "Phong/Don vi",
    ]
    data = [
        "Nguyen Van A", "a@agribank.com.vn", "HQPT01", "001234567890",
        "0982867163", "Đại lý viên", "6900 Phong KHDN",
    ]
    matrix = [header, data]
    cells = [
        CellData(row=r, col=c, text=txt, confidence=1.0, bbox=[])
        for r, row in enumerate(matrix)
        for c, txt in enumerate(row)
    ]
    table = TableData(
        table_index=0,
        num_rows=len(matrix),
        num_cols=len(header),
        cells=cells,
        html="",
        table_kind="",
    )

    users, warnings = map_table_to_users(table)
    assert not any("không tìm thấy cột email" in w for w in warnings)
    assert len(users) == 1
    assert users[0].email.endswith("@agribank.com.vn")
    assert users[0].ipcas_code == "HQPT01"


def test_import_from_docx_without_table_raises(tmp_path):
    from docx import Document

    docx_path = tmp_path / "no_table.docx"
    doc = Document()
    doc.add_paragraph("Chỉ có văn bản, không có bảng.")
    doc.save(str(docx_path))

    with pytest.raises(ValueError):
        import_from_docx(docx_path, "job-docx", "no_table.docx")
