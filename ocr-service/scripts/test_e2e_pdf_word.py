"""
E2E test: PDF text import → export Word → re-import Word.
Chạy: python scripts/test_e2e_pdf_word.py [đường_dẫn_pdf]
"""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document

from app.services.docx_service import export_to_docx, import_from_docx
from app.services.pdf_text_service import import_from_pdf_text
from app.services.user_mapping import map_table_to_users, validate_ocr_result


def main() -> int:
    pdf = Path(
        sys.argv[1]
        if len(sys.argv) > 1
        else r"C:\Users\lamanhzuto2k\Downloads\Telegram Desktop\3526_DT_AgribankBanca (1) (1).pdf"
    )
    if not pdf.exists():
        print(f"KHONG TIM THAY PDF: {pdf}")
        return 1

    print(f"=== PDF text import: {pdf.name} ===")
    result = import_from_pdf_text(pdf, "e2e-test", pdf.name)
    if not result:
        print("FAIL: khong trich duoc bang tu PDF")
        return 1

    total_users = 0
    for page in result.pages:
        for table in page.tables:
            users, _ = map_table_to_users(table)
            total_users += len(users)
            print(f"  Trang {page.page_number}: {len(users)} user, {table.num_rows} dong")

    print(f"  Tong: {total_users} user / {result.total_pages} trang")

    validation = validate_ocr_result(result)
    print(
        f"  Validation: {len(validation.get('errors', []))} loi, "
        f"{len(validation.get('warnings', []))} canh bao"
    )

    print("\n=== Export Word ===")
    out_docx = export_to_docx(result)
    print(f"  Da xuat: {out_docx} ({out_docx.stat().st_size // 1024} KB)")

    doc = Document(str(out_docx))
    print(f"  Bang trong Word: {len(doc.tables)}, dong dau: {len(doc.tables[0].rows)}")

    print("\n=== Re-import Word da xuat ===")
    reimported = import_from_docx(out_docx, "e2e-test", out_docx.name)
    re_users = 0
    for page in reimported.pages:
        for table in page.tables:
            users, _ = map_table_to_users(table)
            re_users += len(users)
    print(f"  Re-import: {re_users} user")

    if re_users < total_users * 0.8:
        print("FAIL: mat user sau vong Word")
        return 1

    # Kiem tra IPCAS QSO*
    qso = 0
    for page in result.pages:
        for table in page.tables:
            for c in table.cells:
                if c.col == 4 and c.text.strip().upper().startswith("QSO"):
                    qso += 1
    print(f"\n  IPCAS QSO*: {qso} o")

    print("\n=== PASS ===")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
