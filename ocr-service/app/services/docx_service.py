"""
DOCX Service — Import dữ liệu bảng từ file Word (.docx) trực tiếp.

Word gốc chứa văn bản thật nên đọc trực tiếp cho độ chính xác gần như tuyệt
đối (không qua OCR). Luồng này song song với import_from_excel: chuyển các
bảng trong tài liệu thành OcrResult (confidence = 1.0) để vào thẳng bước
kiểm tra/tạo lô.
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from app.models.schemas import OcrResult, PageResult, TableData
from app.services.email_reconcile import email_mismatch_with_ipcas, email_needs_review
from app.services.excel_service import _matrix_to_table
from app.services.user_mapping import (
    _normalize_header_key,
    _strip_sso_preamble_rows,
    table_has_user_columns,
)

logger = logging.getLogger(__name__)


def _clean_cell_text(text: str) -> str:
    """Chuẩn hóa text 1 ô: bỏ khoảng trắng thừa, gộp xuống dòng thành space."""
    if not text:
        return ""
    parts = [p.strip() for p in str(text).splitlines()]
    joined = " ".join(p for p in parts if p)
    return " ".join(joined.split())


def _table_to_matrix(table) -> list[list[str]]:
    """Đọc một bảng python-docx thành ma trận chuỗi, bỏ dòng/cột rỗng cuối."""
    rows: list[list[str]] = []
    max_col = 0
    for row in table.rows:
        vals: list[str] = []
        row_non_empty = False
        for c, cell in enumerate(row.cells):
            txt = _clean_cell_text(cell.text)
            vals.append(txt)
            if txt:
                row_non_empty = True
                max_col = max(max_col, c + 1)
        if row_non_empty:
            rows.append(vals)

    if not rows or max_col == 0:
        return []
    return [r[:max_col] + [""] * (max_col - len(r)) for r in rows]


def import_from_docx(
    docx_path: str | Path,
    job_id: str,
    filename: str,
) -> OcrResult:
    """
    Nạp dữ liệu bảng từ file Word (.docx) vào cấu trúc OcrResult.

    Mỗi bảng trong tài liệu được coi là một TableData. Toàn bộ bảng gộp vào
    một "trang" duy nhất (giống một sheet Excel). Ô có confidence = 1.0.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - phụ thuộc runtime
        raise ValueError(
            "Thiếu thư viện python-docx. Cài đặt: pip install python-docx"
        ) from exc

    document = Document(str(docx_path))

    tables: list[TableData] = []
    for tbl in document.tables:
        matrix = _table_to_matrix(tbl)
        if not matrix or not table_has_user_columns(matrix):
            continue
        matrix = _strip_sso_preamble_rows(matrix)
        if not matrix:
            continue
        table = _matrix_to_table(matrix, len(tables))
        if table.cells:
            tables.append(table)

    if not tables:
        raise ValueError(
            "Không tìm thấy bảng danh sách user trong file Word. "
            "Đảm bảo có bảng đúng mẫu SSO 10 cột (có cột Email hoặc User IPCAS)."
        )

    pages = [
        PageResult(
            page_number=1,
            image_path="",
            tables=tables,
            raw_text="",
        )
    ]

    now = datetime.now()
    logger.info(
        "DOCX imported: %s — %d bảng, %d dòng dữ liệu",
        filename,
        len(tables),
        sum(t.num_rows for t in tables),
    )
    return OcrResult(
        job_id=job_id,
        filename=filename,
        total_pages=len(pages),
        pages=pages,
        is_complete=True,
        created_at=now,
        updated_at=now,
    )


# ── Word export styles (khớp palette Excel) ──────────────────────────────

_DOCX_HEADER_FILL = "1F4E79"
_DOCX_ALT_ROW_FILL = "F5F9FC"
_DOCX_LOW_CONF_FILL = "FFF2CC"
_DOCX_EMAIL_WARN_FILL = "FFCDD2"
_DOCX_ROLE_WARN_FILL = "FFF9C4"


def _docx_set_cell_shading(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(shading)


def _docx_style_run(run, *, bold=False, size_pt=10, color=None, italic=False) -> None:
    from docx.shared import Pt, RGBColor

    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _docx_write_cell(
    cell,
    text: str,
    *,
    align_center: bool = False,
    fill: str | None = None,
    bold: bool = False,
    italic: bool = False,
    font_color: tuple[int, int, int] | None = None,
    size_pt: int = 10,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text or "")
    _docx_style_run(
        run,
        bold=bold,
        italic=italic,
        size_pt=size_pt,
        color=font_color,
    )
    if fill:
        _docx_set_cell_shading(cell, fill)


def _docx_configure_landscape(doc) -> None:
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm, Inches

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.2)
    # Độ rộng cột gợi ý (11 cột SSO)
    _DOCX_COL_WIDTHS = (
        Inches(0.35),  # STT
        Inches(1.55),  # Họ tên
        Inches(0.55),  # Mã CN
        Inches(1.35),  # Tên CN
        Inches(0.75),  # IPCAS
        Inches(0.95),  # CCCD
        Inches(1.65),  # Email
        Inches(0.75),  # SĐT
        Inches(0.95),  # Phân quyền
        Inches(0.75),  # Mã LNH
        Inches(1.05),  # Vai trò gợi ý
    )
    return _DOCX_COL_WIDTHS


def export_to_docx(
    result: OcrResult,
    *,
    page_numbers: list[int] | None = None,
    template_id: str | None = None,
) -> Path:
    """
    Xuất kết quả OCR/import ra file Word (.docx) — một bảng SSO mỗi trang.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ValueError(
            "Thiếu thư viện python-docx. Cài đặt: pip install python-docx"
        ) from exc

    from app.config import settings
    from app.services.excel_service import (
        SSO_HEADERS,
        _SSO_GRID_TO_EXCEL_COL_10,
        _SSO_GRID_TO_EXCEL_COL_9,
    )
    from app.services.template_service import (
        export_headers_for,
        get_template_or_default,
    )
    from app.services.user_mapping import (
        _parse_branch_code_digits,
        _parse_department_cell,
        normalize_roles,
    )

    profile = get_template_or_default(template_id)
    headers = export_headers_for(profile) or list(SSO_HEADERS)
    doc_title = (
        profile.export.docx_title
        or profile.name
        or "Danh sách dữ liệu OCR"
    )

    page_filter = set(page_numbers) if page_numbers else None
    pages = result.pages
    if page_filter:
        pages = [p for p in result.pages if p.page_number in page_filter]
    if not pages:
        raise ValueError("Không có trang nào để xuất Word")

    doc = Document()
    col_widths = _docx_configure_landscape(doc)
    threshold = settings.ocr_confidence_threshold

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(doc_title)
    _docx_style_run(tr, bold=True, size_pt=14, color=(0x1F, 0x4E, 0x79))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stem = Path(result.filename).stem or result.job_id
    mr = meta.add_run(
        f"Job: {result.job_id}  ·  File: {stem}  ·  "
        f"{datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    _docx_style_run(mr, size_pt=9, color=(0x66, 0x66, 0x66))
    doc.add_paragraph()

    table_count = 0
    for page in pages:
        sso_tables = [t for t in page.tables if t.table_kind == "sso_agribank"]
        tables_to_export = sso_tables if sso_tables else list(page.tables)
        if not tables_to_export:
            continue
        if len(pages) > 1 or len(tables_to_export) > 1:
            h = doc.add_paragraph()
            hr = h.add_run(f"Trang {page.page_number}")
            _docx_style_run(hr, bold=True, size_pt=12, color=(0x1F, 0x4E, 0x79))

        for table in tables_to_export:
            if not table.cells:
                continue
            table_count += 1
            grid = {(c.row, c.col): (c.text, c.confidence) for c in table.cells}
            is_new_layout = table.num_cols >= 10
            grid_to_excel = (
                _SSO_GRID_TO_EXCEL_COL_10 if is_new_layout else _SSO_GRID_TO_EXCEL_COL_9
            )
            email_grid_col = 6 if is_new_layout else 5
            role_grid_col = 8 if is_new_layout else 7
            ipcas_grid_col = 4 if is_new_layout else 3

            num_cols = len(headers)
            num_rows = max(table.num_rows, 1) + 1
            doc_table = doc.add_table(rows=num_rows, cols=num_cols)
            doc_table.style = "Table Grid"
            doc_table.autofit = False
            for ci, width in enumerate(col_widths):
                if ci < num_cols:
                    doc_table.columns[ci].width = width

            # Header row
            for ci, title_text in enumerate(headers):
                cell = doc_table.rows[0].cells[ci]
                cell.text = title_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _docx_style_run(run, bold=True, size_pt=9, color=(0xFF, 0xFF, 0xFF))
                _docx_set_cell_shading(cell, "1F4E79")

            use_sso_layout = (
                profile.id == "sso-agribank"
                or profile.ocr.sso_enhance
                or table.table_kind == "sso_agribank"
            )

            for r in range(max(table.num_rows, 1)):
                if use_sso_layout and num_cols == len(SSO_HEADERS):
                    dept_text = grid.get((r, 2), ("", 1.0))[0]
                    branch_code = (
                        grid.get((r, 2), ("", 1.0))[0].strip()
                        if is_new_layout
                        else _parse_branch_code_digits(dept_text)
                    )
                    role_text = grid.get((r, role_grid_col), ("", 1.0))[0]
                    role_suggested = (
                        ";".join(normalize_roles(role_text)) if role_text else ""
                    )
                    for excel_c in range(num_cols):
                        cell = doc_table.rows[r + 1].cells[excel_c]
                        if excel_c == 2 and not is_new_layout:
                            cell.text = branch_code
                        elif excel_c == 3 and not is_new_layout:
                            _, _, bn = _parse_department_cell(dept_text)
                            cell.text = bn or dept_text
                        elif excel_c == 10:
                            cell.text = role_suggested
                        else:
                            grid_c = next(
                                (g for g, e in grid_to_excel.items() if e == excel_c),
                                None,
                            )
                            text_val = (
                                grid.get((r, grid_c), ("", 1.0))[0]
                                if grid_c is not None
                                else ""
                            )
                            cell.text = text_val
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                _docx_style_run(run, size_pt=8)
                else:
                    for excel_c in range(num_cols):
                        cell = doc_table.rows[r + 1].cells[excel_c]
                        text_val, conf = grid.get((r, excel_c), ("", 1.0))
                        cell.text = text_val
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                _docx_style_run(run, size_pt=8)
                                if conf < threshold:
                                    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

            doc.add_paragraph()

    if table_count == 0:
        raise ValueError("Không có bảng dữ liệu để xuất Word")

    from app.config import settings as _settings

    base = result.filename.rsplit(".", 1)[0]
    if page_filter:
        page_tag = "-".join(str(p) for p in sorted(page_filter))
        filename = f"{result.job_id}_{base}_trang-{page_tag}.docx"
    else:
        filename = f"{result.job_id}_{base}.docx"
    export_file = _settings.export_path / filename
    doc.save(str(export_file))
    logger.info("Word exported: %s", export_file)
    return export_file
